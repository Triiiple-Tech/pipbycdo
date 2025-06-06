from backend.agents.base_agent import BaseAgent
from backend.app.schemas import AppState, EstimateItem
import logging
import json
from typing import List, Dict, Any
from datetime import datetime, timezone
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from fpdf import FPDF

logger = logging.getLogger(__name__)

class ExporterAgent(BaseAgent):
    """
    Agent responsible for exporting estimate data to various formats.
    Supports JSON, DOCX, PDF, and XLSX export formats with comprehensive formatting.
    """
    
    def __init__(self):
        super().__init__("exporter")
    
    def process(self, state: AppState) -> AppState:
        """Export estimate data to the requested format."""
        self.log_interaction(state, "Starting export process", 
                           "Exporter Agent preparing estimate data for export")
        
        # Validate required data
        if not state.estimate:
            error_msg = "Export failed: Missing estimate data to export"
            state.error = error_msg
            self.log_interaction(state, "Export failed due to missing data", error_msg, level="error")
            return state
        
        # Determine export format
        export_format = self._get_export_format(state)
        
        self.log_interaction(state, f"Exporting to {export_format.upper()} format", 
                           f"Processing {len(state.estimate)} estimate items for export")
        
        try:
            # Perform the export based on format
            if export_format == "json":
                self._export_json(state)
            elif export_format == "docx":
                self._export_docx(state)
            elif export_format == "pdf":
                self._export_pdf(state)
            elif export_format == "xlsx":
                self._export_xlsx(state)
            else:
                error_msg = f"Export format '{export_format}' is not supported"
                state.error = error_msg
                self.log_interaction(state, f"Unsupported export format: {export_format}", 
                                   error_msg, level="error")
                return state
            
            self.log_interaction(state, "Export process complete", 
                               f"Successfully exported estimate to {export_format.upper()} format")
            
            # Set export message for backward compatibility
            state.export = f"Exported estimate with {len(state.estimate)} items."
            
        except Exception as e:
            error_msg = f"Export failed during file generation: {str(e)}"
            state.error = error_msg
            self.log_interaction(state, "Export file generation error", error_msg, level="error")
            # Clear any partial export data
            state.exported_file_content = None
            state.exported_file_name = None
            state.exported_content_type = None
        
        return state
    
    def _get_export_format(self, state: AppState) -> str:
        """Determine the export format from state configuration."""
        if state.export_options and state.export_options.get("format"):
            return state.export_options["format"].lower()
        return "json"  # Default format
    
    def _get_project_name(self, state: AppState) -> str:
        """Extract project name from state metadata or query."""
        if state.metadata and state.metadata.get("project_name") != "N/A":
            return state.metadata["project_name"]
        elif state.query:
            return state.query[:50] + "..." if len(state.query) > 50 else state.query
        return "Construction Project"
    
    def _calculate_grand_total(self, estimate_items: List[EstimateItem]) -> float:
        """Calculate the grand total from estimate items."""
        return sum(item.total for item in estimate_items)
    
    def _export_json(self, state: AppState) -> None:
        """Export estimate data to JSON format."""
        estimate_items = state.estimate
        grand_total = self._calculate_grand_total(estimate_items)
        project_name = self._get_project_name(state)
        
        export_data: Dict[str, Any] = {
            "project_name": project_name,
            "export_date": datetime.now(timezone.utc).isoformat(),
            "grand_total": grand_total,
            "currency": "USD",
            "item_count": len(estimate_items),
            "items": [item.model_dump() for item in estimate_items],
            "metadata": {
                "exported_by": "PipByCDO System",
                "version": "1.0",
                "export_format": "json"
            }
        }
        
        state.exported_file_content = json.dumps(export_data, indent=2).encode('utf-8')
        state.exported_file_name = f"estimate_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
        state.exported_content_type = "application/json"
        
        self.log_interaction(state, f"JSON export prepared: {state.exported_file_name}", 
                           f"Exported {len(estimate_items)} items with total ${grand_total:,.2f}")
    
    def _export_docx(self, state: AppState) -> None:
        """Export estimate data to Word document format."""
        estimate_items = state.estimate
        grand_total = self._calculate_grand_total(estimate_items)
        project_name = self._get_project_name(state)
        
        document = Document()
        
        # Add title and header information
        title = document.add_heading('Construction Estimate', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        document.add_paragraph(f"Project: {project_name}")
        document.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
        document.add_paragraph(f"Total Items: {len(estimate_items)}")
        document.add_paragraph(f"Grand Total: ${grand_total:,.2f}")
        document.add_paragraph()
        
        if estimate_items:
            # Create table with headers
            table = document.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'CSI Division'
            hdr_cells[3].text = 'Qty'
            hdr_cells[4].text = 'Unit'
            hdr_cells[5].text = 'Unit Price'
            hdr_cells[6].text = 'Total'
            
            # Add data rows
            for item in estimate_items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.item)
                row_cells[1].text = str(item.description or "")
                row_cells[2].text = str(item.csi_division or "")
                row_cells[3].text = str(item.qty)
                row_cells[4].text = str(item.unit)
                row_cells[5].text = f"${item.unit_price:,.2f}"
                row_cells[6].text = f"${item.total:,.2f}"
        else:
            document.add_paragraph("No items in estimate.")
        
        # Add footer
        document.add_paragraph()
        footer_p = document.add_paragraph("Generated by PipByCDO Construction Estimating System")
        footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        output = BytesIO()
        document.save(output)
        state.exported_file_content = output.getvalue()
        state.exported_file_name = f"estimate_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
        state.exported_content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        self.log_interaction(state, f"DOCX export prepared: {state.exported_file_name}", 
                           f"Exported {len(estimate_items)} items to Word document")
    
    def _export_pdf(self, state: AppState) -> None:
        """Export estimate data to PDF format."""
        estimate_items = state.estimate
        grand_total = self._calculate_grand_total(estimate_items)
        project_name = self._get_project_name(state)
        
        pdf = FPDF()
        pdf.add_page()
        
        # Title
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Construction Estimate", 0, 1, "C")
        pdf.ln(5)
        
        # Header information
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 8, f"Project: {project_name}", 0, 1)
        pdf.cell(0, 8, f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}", 0, 1)
        pdf.cell(0, 8, f"Total Items: {len(estimate_items)}", 0, 1)
        pdf.cell(0, 8, f"Grand Total: ${grand_total:,.2f}", 0, 1)
        pdf.ln(10)
        
        if estimate_items:
            # Table headers
            pdf.set_font("Arial", "B", 9)
            col_widths = [20, 50, 25, 15, 15, 25, 25]
            headers = ['Item', 'Description', 'CSI Division', 'Qty', 'Unit', 'Unit Price', 'Total']
            
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 8, header, 1, 0, "C")
            pdf.ln()
            
            # Table data
            pdf.set_font("Arial", "", 8)
            for item in estimate_items:
                # Handle long descriptions by truncating
                description = str(item.description or "")
                if len(description) > 30:
                    description = description[:27] + "..."
                
                pdf.cell(col_widths[0], 6, str(item.item), 1, 0, "C")
                pdf.cell(col_widths[1], 6, description, 1, 0, "L")
                pdf.cell(col_widths[2], 6, str(item.csi_division or ""), 1, 0, "C")
                pdf.cell(col_widths[3], 6, str(item.qty), 1, 0, "R")
                pdf.cell(col_widths[4], 6, str(item.unit), 1, 0, "C")
                pdf.cell(col_widths[5], 6, f"${item.unit_price:,.2f}", 1, 0, "R")
                pdf.cell(col_widths[6], 6, f"${item.total:,.2f}", 1, 0, "R")
                pdf.ln()
        else:
            pdf.cell(0, 10, "No items in estimate.", 0, 1)
        
        # Footer
        pdf.ln(10)
        pdf.set_font("Arial", "I", 8)
        pdf.cell(0, 5, "Generated by PipByCDO Construction Estimating System", 0, 1, "C")
        
        state.exported_file_content = pdf.output(dest='B')
        state.exported_file_name = f"estimate_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
        state.exported_content_type = "application/pdf"
        
        self.log_interaction(state, f"PDF export prepared: {state.exported_file_name}", 
                           f"Exported {len(estimate_items)} items to PDF document")
    
    def _export_xlsx(self, state: AppState) -> None:
        """Export estimate data to Excel spreadsheet format."""
        estimate_items = state.estimate
        grand_total = self._calculate_grand_total(estimate_items)
        project_name = self._get_project_name(state)
        
        workbook = Workbook()
        sheet = workbook.active
        if sheet is None:
            raise Exception("Failed to get active sheet from new workbook")
        
        sheet.title = "Estimate"
        
        # Header information
        sheet.append([f"Project: {project_name}"])
        sheet.append([f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"])
        sheet.append([f"Total Items: {len(estimate_items)}"])
        sheet.append([f"Grand Total:", grand_total])
        sheet['B4'].number_format = '$#,##0.00'
        sheet.append([])  # Empty row
        
        if estimate_items:
            # Column headers
            headers = ['Item', 'Description', 'CSI Division', 'Qty', 'Unit', 'Unit Price', 'Total']
            sheet.append(headers)
            
            # Data rows
            for item in estimate_items:
                sheet.append([
                    item.item,
                    item.description or "",
                    item.csi_division or "",
                    item.qty,
                    item.unit,
                    item.unit_price,
                    item.total
                ])
            
            # Format currency columns
            for row_idx in range(7, sheet.max_row + 1):  # Starting from first data row
                sheet[f'F{row_idx}'].number_format = '$#,##0.00'  # Unit Price
                sheet[f'G{row_idx}'].number_format = '$#,##0.00'  # Total
        else:
            sheet.append(["No items in estimate."])
        
        # Add footer
        sheet.append([])
        sheet.append(["Generated by PipByCDO Construction Estimating System"])
        
        output = BytesIO()
        workbook.save(output)
        state.exported_file_content = output.getvalue()
        state.exported_file_name = f"estimate_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.xlsx"
        state.exported_content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        self.log_interaction(state, f"XLSX export prepared: {state.exported_file_name}", 
                           f"Exported {len(estimate_items)} items to Excel spreadsheet")


# Create singleton instance
exporter_agent = ExporterAgent()

# Backward compatibility function
def handle(state_dict: Dict[str, Any]) -> Dict[str, Any]:
    """Legacy handle function for backward compatibility."""
    return exporter_agent.handle(state_dict)
