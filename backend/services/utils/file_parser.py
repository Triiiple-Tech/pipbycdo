from typing import Dict, List, Optional, Union, Any
from backend.app.schemas import File, AppState
from backend.services.utils.state import log_agent_interaction
import io
import logging

# Import optional dependencies
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

logger = logging.getLogger(__name__)

class MultimodalFileParser:
    """
    Centralized file parser that handles multiple file types and formats.
    Supports text, PDF, DOCX, XLSX, and image files with OCR capabilities.
    """
    
    SUPPORTED_TYPES = {
        'text/plain': 'text',
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
        'image/png': 'image',
        'image/jpeg': 'image',
        'image/jpg': 'image',
        'image/gif': 'image',
        'image/bmp': 'image',
        'image/tiff': 'image'
    }
    
    def __init__(self):
        self.parsers = {
            'text': self._parse_text,
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'xlsx': self._parse_xlsx,
            'image': self._parse_image
        }
    
    def parse_file(self, file: File, state_dict: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Parse a single file and return parsing results.
        
        Returns:
            Dict with keys: 'content', 'status', 'error', 'metadata'
        """
        if not file.data:
            return {
                'content': '',
                'status': 'error',
                'error': 'No file data provided',
                'metadata': {'parsed_pages': 0, 'parsed_chars': 0}
            }
        
        # Determine file type from metadata or filename
        content_type = file.metadata.get('content_type')
        if not content_type:
            content_type = self._guess_content_type(file.filename)
        
        file_type = self.SUPPORTED_TYPES.get(content_type)
        if not file_type:
            return {
                'content': f'[Unsupported file type: {content_type}]',
                'status': 'error',
                'error': f'Unsupported content type: {content_type}',
                'metadata': {'content_type': content_type}
            }
        
        # Log parsing attempt
        if state_dict:
            log_agent_interaction(
                state_dict, 
                "file_parser", 
                f"Parsing {file.filename} ({file_type})", 
                f"Starting to parse {file.filename} as {file_type}"
            )
        
        try:
            parser = self.parsers.get(file_type)
            if not parser:
                return {
                    'content': f'[No parser available for {file_type}]',
                    'status': 'error',
                    'error': f'No parser implemented for {file_type}',
                    'metadata': {'file_type': file_type}
                }
            
            result = parser(file.data, file.filename)
            
            # Update file object with parsed content
            file.content = result['content']
            file.status = result['status']
            
            # Log success
            if state_dict and result['status'] == 'parsed':
                log_agent_interaction(
                    state_dict,
                    "file_parser",
                    f"Successfully parsed {file.filename}",
                    f"Extracted {len(result['content'])} characters from {file.filename}"
                )
            
            return result
            
        except Exception as e:
            error_msg = f"Error parsing {file.filename}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            
            if state_dict:
                log_agent_interaction(
                    state_dict,
                    "file_parser",
                    f"Failed to parse {file.filename}",
                    error_msg,
                    level="error"
                )
            
            return {
                'content': f'[Error: {str(e)}]',
                'status': 'error',
                'error': str(e),
                'metadata': {'file_type': file_type}
            }
    
    def parse_files(self, files: List[File], state_dict: Optional[Dict[str, Any]] = None) -> Dict[str, Dict[str, Any]]:
        """
        Parse multiple files and return a dictionary of results.
        
        Returns:
            Dict[filename, parsing_result]
        """
        results = {}
        
        for file in files:
            results[file.filename] = self.parse_file(file, state_dict)
        
        return results
    
    def _guess_content_type(self, filename: str) -> str:
        """Guess content type from file extension."""
        if not filename:
            return 'application/octet-stream'
        
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        ext_mapping = {
            'txt': 'text/plain',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'xls': 'application/vnd.ms-excel',
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'tiff': 'image/tiff'
        }
        
        return ext_mapping.get(ext, 'application/octet-stream')
    
    def _parse_text(self, data: bytes, filename: str) -> Dict[str, Any]:
        """Parse plain text files."""
        try:
            content = data.decode('utf-8')
            return {
                'content': content,
                'status': 'parsed',
                'error': None,
                'metadata': {
                    'parsed_chars': len(content),
                    'encoding': 'utf-8'
                }
            }
        except UnicodeDecodeError:
            try:
                content = data.decode('latin-1')
                return {
                    'content': content,
                    'status': 'parsed',
                    'error': None,
                    'metadata': {
                        'parsed_chars': len(content),
                        'encoding': 'latin-1'
                    }
                }
            except Exception as e:
                return {
                    'content': '[Error: Could not decode text file]',
                    'status': 'error',
                    'error': f'Text decoding error: {str(e)}',
                    'metadata': {}
                }
    
    def _parse_pdf(self, data: bytes, filename: str) -> Dict[str, Any]:
        """Parse PDF files."""
        if not PyPDF2:
            return {
                'content': '[Skipped: PyPDF2 library not available]',
                'status': 'error',
                'error': 'PyPDF2 library not available',
                'metadata': {'missing_dependency': 'PyPDF2'}
            }
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(data))
            text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num in range(page_count):
                page_text = pdf_reader.pages[page_num].extract_text()
                text += page_text + "\n"
            
            return {
                'content': text.strip(),
                'status': 'parsed',
                'error': None,
                'metadata': {
                    'parsed_pages': page_count,
                    'parsed_chars': len(text),
                    'parser': 'PyPDF2'
                }
            }
        except Exception as e:
            return {
                'content': f'[Error extracting PDF content: {str(e)}]',
                'status': 'error',
                'error': str(e),
                'metadata': {'parser': 'PyPDF2'}
            }
    
    def _parse_docx(self, data: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX files."""
        if not docx:
            return {
                'content': '[Skipped: python-docx library not available]',
                'status': 'error',
                'error': 'python-docx library not available',
                'metadata': {'missing_dependency': 'python-docx'}
            }
        
        try:
            document = docx.Document(io.BytesIO(data))
            paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            # Extract table content as well
            tables_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            if tables_text:
                text += "\n\nTables:\n" + "\n".join(tables_text)
            
            return {
                'content': text,
                'status': 'parsed',
                'error': None,
                'metadata': {
                    'parsed_paragraphs': len(paragraphs),
                    'parsed_tables': len(document.tables),
                    'parsed_chars': len(text),
                    'parser': 'python-docx'
                }
            }
        except Exception as e:
            return {
                'content': f'[Error extracting DOCX content: {str(e)}]',
                'status': 'error',
                'error': str(e),
                'metadata': {'parser': 'python-docx'}
            }
    
    def _parse_xlsx(self, data: bytes, filename: str) -> Dict[str, Any]:
        """Parse XLSX files."""
        if not load_workbook:
            return {
                'content': '[Skipped: openpyxl library not available]',
                'status': 'error',
                'error': 'openpyxl library not available',
                'metadata': {'missing_dependency': 'openpyxl'}
            }
        
        try:
            workbook = load_workbook(io.BytesIO(data), read_only=True)
            text_parts = []
            total_rows = 0
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                max_row = worksheet.max_row or 0
                max_col = worksheet.max_column or 0
                
                # Limit to prevent huge outputs
                max_row = min(max_row, 1000)
                max_col = min(max_col, 50)
                
                for row_num in range(1, max_row + 1):
                    row_data = []
                    has_data = False
                    
                    for col_num in range(1, max_col + 1):
                        cell_value = worksheet.cell(row_num, col_num).value
                        if cell_value is not None:
                            has_data = True
                            # Handle different data types
                            if isinstance(cell_value, (int, float)):
                                cell_str = str(cell_value)
                            else:
                                from datetime import date, datetime
                                if isinstance(cell_value, (date, datetime)):
                                    try:
                                        cell_str = cell_value.strftime('%Y-%m-%d %H:%M:%S')
                                    except AttributeError:
                                        cell_str = str(cell_value)
                                else:
                                    cell_str = str(cell_value)
                            row_data.append(cell_str)
                        else:
                            row_data.append("")
                    
                    if has_data:
                        # Remove trailing empty cells
                        while row_data and row_data[-1] == "":
                            row_data.pop()
                        
                        if row_data:
                            text_parts.append(" | ".join(row_data))
                            total_rows += 1
            
            workbook.close()
            text = "\n".join(text_parts)
            
            return {
                'content': text,
                'status': 'parsed',
                'error': None,
                'metadata': {
                    'parsed_sheets': len(workbook.sheetnames),
                    'parsed_rows': total_rows,
                    'parsed_chars': len(text),
                    'parser': 'openpyxl'
                }
            }
        except Exception as e:
            return {
                'content': f'[Error extracting XLSX content: {str(e)}]',
                'status': 'error',
                'error': str(e),
                'metadata': {'parser': 'openpyxl'}
            }
    
    def _parse_image(self, data: bytes, filename: str) -> Dict[str, Any]:
        """Parse image files using OCR."""
        if not Image or not pytesseract:
            return {
                'content': '[Skipped: Pillow or pytesseract library not available]',
                'status': 'error',
                'error': 'Pillow or pytesseract library not available',
                'metadata': {'missing_dependency': 'Pillow or pytesseract'}
            }
        
        try:
            image = Image.open(io.BytesIO(data))
            
            # Get image info
            width, height = image.size
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            return {
                'content': text.strip(),
                'status': 'parsed',
                'error': None,
                'metadata': {
                    'image_width': width,
                    'image_height': height,
                    'image_mode': image.mode,
                    'parsed_chars': len(text),
                    'parser': 'pytesseract'
                }
            }
        except Exception as e:
            return {
                'content': f'[Error processing image with OCR: {str(e)}]',
                'status': 'error',
                'error': str(e),
                'metadata': {'parser': 'pytesseract'}
            }

# Singleton instance for easy importing
file_parser = MultimodalFileParser()

def parse_file(file: File, state_dict: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Convenience function to parse a single file."""
    return file_parser.parse_file(file, state_dict)

def parse_files(files: List[File], state_dict: Optional[Dict[str, Any]] = None) -> Dict[str, Dict[str, Any]]:
    """Convenience function to parse multiple files."""
    return file_parser.parse_files(files, state_dict)
